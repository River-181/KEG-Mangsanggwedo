from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[4]
TEMPLATE_PATH = ROOT / "05_제출/20260406_③_2026_KIT_바이브코딩_공모전_팀명_개인은_이름_AI리포트.docx"
CONTENT_PATH = ROOT / "05_제출/ai-report-final.md"
OUTPUT_DOCX_PATH = ROOT / "05_제출/20260413_③_2026_KIT_바이브코딩_공모전_망상궤도_AI리포트.docx"
ASSET_SEARCH_DIRS = [
    ROOT / "assets/excalidraw/png",
    ROOT / "assets/screenshots",
]

TEAM_NAME = "망상궤도"
MEMBERS = "이승보(CEO), 김주용(COO)"
SUBMISSION_DATE = "2026-04-13"
REPRESENTATIVE_PHONE = "[ 대표자 전화번호 ]"

QUESTION_HEADING_RE = re.compile(r"^Q([1-6])\.")
IMAGE_LINE_RE = re.compile(r"!\[\[([^\]]+\.(?:png|PNG))\]\]")
SPECIAL_IMAGE_RE = re.compile(r"\{\{[^}\n]*?([A-Za-z0-9_\-./가-힣 ]+\.(?:png|PNG))[^}\n]*\}\}")
TOP_LEVEL_HEADING_RE = re.compile(r"^## (.+)$", re.MULTILINE)
PROJECT_NAME_RE = re.compile(r"\| 프로젝트명 \| \*\*(.+?)\*\* \|")


@dataclass
class BuildStats:
    image_count: int = 0


def main() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Missing template: {TEMPLATE_PATH}")
    if not CONTENT_PATH.exists():
        raise FileNotFoundError(f"Missing content source: {CONTENT_PATH}")

    OUTPUT_DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE_PATH, OUTPUT_DOCX_PATH)

    markdown = CONTENT_PATH.read_text(encoding="utf-8")
    question_sections = extract_question_sections(markdown)
    project_name = extract_project_name(markdown)

    missing_questions = [f"Q{i}" for i in range(1, 7) if str(i) not in question_sections]
    if missing_questions:
        raise ValueError(f"Missing question sections: {', '.join(missing_questions)}")

    doc = Document(str(OUTPUT_DOCX_PATH))
    apply_core_properties(doc, project_name)
    apply_header_metadata(doc)
    apply_cover_table(doc, project_name)

    stats = BuildStats()
    for index, question_number in enumerate(range(1, 7), start=1):
        cell = doc.tables[index].cell(0, 0)
        clear_cell(cell)
        write_markdown_into_cell(cell, question_sections[str(question_number)], stats)

    doc.save(str(OUTPUT_DOCX_PATH))
    verify_output_docx(question_sections)

    print(f"DOCX: {OUTPUT_DOCX_PATH}")
    print(f"Embedded images: {stats.image_count}")
    print("Questions: Q1-Q6")


def extract_question_sections(markdown: str) -> dict[str, str]:
    headings = list(TOP_LEVEL_HEADING_RE.finditer(markdown))
    sections: dict[str, str] = {}
    for index, match in enumerate(headings):
        heading_text = match.group(1).strip()
        question_match = QUESTION_HEADING_RE.match(heading_text)
        if not question_match:
            continue
        start = match.end()
        end = headings[index + 1].start() if index + 1 < len(headings) else len(markdown)
        sections[question_match.group(1)] = markdown[start:end].strip()
    return sections


def extract_project_name(markdown: str) -> str:
    match = PROJECT_NAME_RE.search(markdown)
    if match:
        return strip_markdown(match.group(1).strip())
    return "HagentOS - AI 에이전트 기반 학원 운영 플랫폼"


def apply_core_properties(doc: Document, project_name: str) -> None:
    core = doc.core_properties
    core.title = project_name
    core.subject = f"{TEAM_NAME} AI 리포트"
    core.author = TEAM_NAME
    core.comments = (
        f"팀명: {TEAM_NAME} | 구성원: {MEMBERS} | 제출일: {SUBMISSION_DATE} | "
        f"대표자 전화번호: {REPRESENTATIVE_PHONE}"
    )
    core.keywords = "AI report, HagentOS, 망상궤도"


def apply_header_metadata(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        extra = header.add_paragraph()
        extra.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = extra.add_run(
            f"팀명: {TEAM_NAME} | 구성원: {MEMBERS} | 제출일: {SUBMISSION_DATE} | "
            f"대표자 전화번호: {REPRESENTATIVE_PHONE}"
        )
        run.font.size = Pt(8.5)


def apply_cover_table(doc: Document, project_name: str) -> None:
    cover = doc.tables[0]
    cover.cell(0, 1).text = TEAM_NAME
    cover.cell(0, 3).text = REPRESENTATIVE_PHONE
    project_cell = cover.cell(1, 1)
    project_cell.text = project_name
    paragraph = project_cell.add_paragraph()
    paragraph.add_run(f"구성원: {MEMBERS}")
    paragraph = project_cell.add_paragraph()
    paragraph.add_run(f"제출일: {SUBMISSION_DATE}")


def clear_cell(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag.endswith("}tcPr"):
            continue
        tc.remove(child)
    tc.append(OxmlElement("w:p"))


def write_markdown_into_cell(cell, markdown: str, stats: BuildStats) -> None:
    state = CellWriterState(cell)
    lines = markdown.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            if index < len(lines):
                index += 1
            add_code_block(state, code_lines)
            continue

        image_path = resolve_image_line(stripped)
        if image_path:
            add_image(state, image_path, stats)
            index += 1
            continue

        if is_table_start(lines, index):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(state, table_lines)
            continue

        if stripped.startswith("### "):
            add_heading(state, stripped[4:].strip(), level=3)
            index += 1
            continue

        if stripped.startswith("#### "):
            add_heading(state, stripped[5:].strip(), level=4)
            index += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            add_quote(state, " ".join(quote_lines))
            continue

        if is_caption_line(stripped):
            add_caption(state, stripped)
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            while index < len(lines) and re.match(r"^\d+\.\s+", lines[index].strip()):
                add_list_item(state, re.sub(r"^\d+\.\s+", "", lines[index].strip()), ordered=True)
                index += 1
            continue

        if stripped.startswith("- "):
            while index < len(lines) and lines[index].strip().startswith("- "):
                add_list_item(state, lines[index].strip()[2:].strip(), ordered=False)
                index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                index += 1
                break
            if (
                candidate == "---"
                or candidate.startswith(("### ", "#### ", ">"))
                or candidate.startswith("- ")
                or re.match(r"^\d+\.\s+", candidate)
                or candidate.startswith("```")
                or candidate.startswith("|")
                or resolve_image_line(candidate)
                or is_caption_line(candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        add_paragraph(state, " ".join(paragraph_lines))


class CellWriterState:
    def __init__(self, cell) -> None:
        self.cell = cell
        self.first_paragraph_available = True

    def paragraph(self, style: str | None = None):
        if self.first_paragraph_available and self.cell.paragraphs:
            paragraph = self.cell.paragraphs[0]
            self.first_paragraph_available = False
            apply_paragraph_style(paragraph, style)
            return paragraph
        paragraph = self.cell.add_paragraph()
        apply_paragraph_style(paragraph, style)
        return paragraph


def apply_paragraph_style(paragraph, style: str | None) -> None:
    if not style:
        return
    try:
        paragraph.style = style
    except KeyError:
        return


def add_heading(state: CellWriterState, text: str, level: int) -> None:
    style = "Heading 3" if level == 3 else "Heading 4"
    paragraph = state.paragraph(style=style)
    paragraph.add_run(strip_markdown(text))


def add_quote(state: CellWriterState, text: str) -> None:
    paragraph = state.paragraph(style="Intense Quote")
    add_inline_runs(paragraph, text)


def add_caption(state: CellWriterState, text: str) -> None:
    paragraph = state.paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(strip_wrapping_asterisks(text))
    run.italic = True
    run.font.size = Pt(9)


def add_list_item(state: CellWriterState, text: str, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    paragraph = state.paragraph(style=style)
    add_inline_runs(paragraph, text)


def add_paragraph(state: CellWriterState, text: str) -> None:
    paragraph = state.paragraph()
    add_inline_runs(paragraph, text)


def add_code_block(state: CellWriterState, lines: list[str]) -> None:
    for line in lines or [""]:
        paragraph = state.paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_markdown_table(state: CellWriterState, lines: list[str]) -> None:
    rows = [parse_markdown_table_row(line) for line in lines]
    if len(rows) < 2:
        for row in rows:
            add_paragraph(state, " | ".join(row))
        return

    header = rows[0]
    body_rows = rows[2:] if is_alignment_row(rows[1]) else rows[1:]
    table = state.cell.add_table(rows=1 + len(body_rows), cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    for column_index, value in enumerate(header):
        paragraph = table.rows[0].cells[column_index].paragraphs[0]
        run = paragraph.add_run(strip_markdown(value))
        run.bold = True

    for row_index, values in enumerate(body_rows, start=1):
        for column_index, value in enumerate(values):
            table.rows[row_index].cells[column_index].text = strip_markdown(value)

    state.first_paragraph_available = False


def add_image(state: CellWriterState, image_path: Path, stats: BuildStats) -> None:
    paragraph = state.paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))
    stats.image_count += 1


def parse_markdown_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return cells


def is_alignment_row(row: Iterable[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in row)


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    first = lines[index].strip()
    second = lines[index + 1].strip()
    return first.startswith("|") and second.startswith("|") and all(
        re.fullmatch(r":?-{3,}:?", cell.strip() or "")
        for cell in second.strip("|").split("|")
    )


def is_caption_line(line: str) -> bool:
    stripped = line.strip()
    return (
        stripped.startswith("*")
        and stripped.endswith("*")
        and ("▲" in stripped or stripped.lower().startswith("*q"))
    )


def resolve_image_line(line: str) -> Path | None:
    for pattern in (IMAGE_LINE_RE, SPECIAL_IMAGE_RE):
        match = pattern.search(line)
        if match:
            return resolve_image_path(match.group(1))
    return None


def resolve_image_path(reference: str) -> Path | None:
    reference_path = Path(reference)
    candidates: list[Path] = []

    if reference_path.is_absolute():
        candidates.append(reference_path)
    else:
        candidates.append(ROOT / reference_path)
        for asset_dir in ASSET_SEARCH_DIRS:
            candidates.append(asset_dir / reference_path.name)

    seen: set[str] = set()
    for candidate in candidates:
        normalized = str(candidate)
        if normalized in seen:
            continue
        seen.add(normalized)
        if candidate.exists():
            return candidate
    return None


def add_inline_runs(paragraph, text: str) -> None:
    cleaned = text.replace("\\|", "|")
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))", cleaned)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            continue
        link_match = re.fullmatch(r"\[(.*?)\]\((.*?)\)", part)
        if link_match:
            paragraph.add_run(link_match.group(1))
            continue
        paragraph.add_run(strip_wrapping_asterisks(part))


def strip_wrapping_asterisks(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("*") and stripped.endswith("*") and len(stripped) >= 2:
        return stripped[1:-1]
    return text


def strip_markdown(text: str) -> str:
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return strip_wrapping_asterisks(text).strip()


def verify_output_docx(question_sections: dict[str, str]) -> None:
    if not OUTPUT_DOCX_PATH.exists():
        raise FileNotFoundError(f"Output docx not found: {OUTPUT_DOCX_PATH}")
    size = OUTPUT_DOCX_PATH.stat().st_size
    if size <= 10 * 1024:
        raise ValueError(f"Output docx is too small: {size} bytes")

    doc = Document(str(OUTPUT_DOCX_PATH))
    for index, question_number in enumerate(range(1, 7), start=1):
        cell_text = doc.tables[index].cell(0, 0).text.strip()
        if not cell_text:
            raise ValueError(f"Question Q{question_number} answer is empty in output docx")
        if not question_sections[str(question_number)].strip():
            raise ValueError(f"Question Q{question_number} source content is empty")

    print(f"Verified size: {size} bytes")


if __name__ == "__main__":
    main()

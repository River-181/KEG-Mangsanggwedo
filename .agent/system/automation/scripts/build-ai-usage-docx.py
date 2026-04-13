"""Fill ONLY Q4/Q5/Q6 (AI 활용 전략) into a standalone docx.

Uses the existing `build-ai-report-docx.py` functions via importlib so we don't
duplicate the markdown-to-docx rendering logic.
"""
from __future__ import annotations

import importlib.util
import shutil
import sys
from pathlib import Path

from docx import Document


SCRIPT_DIR = Path(__file__).resolve().parent
SHARED_SCRIPT = SCRIPT_DIR / "build-ai-report-docx.py"


def _load_shared():
    spec = importlib.util.spec_from_file_location("build_ai_report_docx", SHARED_SCRIPT)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules["build_ai_report_docx"] = module  # required for @dataclass resolution
    spec.loader.exec_module(module)
    return module


_shared = _load_shared()


ROOT = Path(__file__).resolve().parents[4]
TEMPLATE_PATH = ROOT / "assets/originals/③_2026_KIT_바이브코딩_공모전_팀명_개인은_이름_AI리포트.docx"
CONTENT_PATH = ROOT / "05_제출/ai-report-final.md"
OUTPUT_DOCX_PATH = (
    ROOT / "05_제출/20260413_③_2026_KIT_바이브코딩_공모전_망상궤도_AI리포트_AI활용전략.docx"
)

TARGET_QUESTIONS = (4, 5, 6)


def main() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Missing template: {TEMPLATE_PATH}")
    if not CONTENT_PATH.exists():
        raise FileNotFoundError(f"Missing content: {CONTENT_PATH}")

    OUTPUT_DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TEMPLATE_PATH, OUTPUT_DOCX_PATH)

    markdown = CONTENT_PATH.read_text(encoding="utf-8")
    sections = _shared.extract_question_sections(markdown)

    missing = [f"Q{n}" for n in TARGET_QUESTIONS if str(n) not in sections]
    if missing:
        raise ValueError(f"Missing sections: {', '.join(missing)}")

    doc = Document(str(OUTPUT_DOCX_PATH))

    stats = _shared.BuildStats()
    for question_number in TARGET_QUESTIONS:
        cell = doc.tables[question_number].cell(0, 0)
        _shared.clear_cell(cell)
        _shared.write_markdown_into_cell(cell, sections[str(question_number)], stats)

    doc.save(str(OUTPUT_DOCX_PATH))

    size = OUTPUT_DOCX_PATH.stat().st_size
    if size <= 10 * 1024:
        raise ValueError(f"Output docx too small: {size} bytes")

    verify_doc = Document(str(OUTPUT_DOCX_PATH))
    for question_number in TARGET_QUESTIONS:
        cell_text = verify_doc.tables[question_number].cell(0, 0).text.strip()
        if not cell_text:
            raise ValueError(f"Q{question_number} cell ended up empty")

    print(f"DOCX: {OUTPUT_DOCX_PATH}")
    print(f"Size: {size} bytes")
    print(f"Filled questions: {', '.join(f'Q{n}' for n in TARGET_QUESTIONS)}")
    print(f"Embedded images: {stats.image_count}")


if __name__ == "__main__":
    main()
